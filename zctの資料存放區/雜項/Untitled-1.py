import json
import pandas as pd
from docx import Document
from docx.shared import Inches

# 載入JSON數據
with open('data.json', 'r', encoding='utf-8') as file:
    data = json.load(file)

# 將JSON數據轉換為DataFrame
df = pd.DataFrame(data)

# 創建一個Word文件
doc = Document()
doc.add_heading('JSON Data Table', level=1)

# 將DataFrame數據添加到Word表格中
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

# 添加表頭
hdr_cells = table.rows[0].cells
for i, column in enumerate(df.columns):
    hdr_cells[i].text = str(column)

# 添加表格內容
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# 儲存Word文件
doc.save('output.docx')
