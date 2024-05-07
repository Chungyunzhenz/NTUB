import json
from docx import Document

# 載入 JSON 檔案
with open("C:/Users/zct/Documents/GitHub/NTUB/zctの資料存放區/output.json", 'r', encoding='utf-8') as file:
    data = json.load(file)

# 創建 Word 文件
doc = Document()

# 添加標題
doc.add_heading('JSON Data to Table', level=1)

# 檢查資料內容
if data and isinstance(data, list):
    # 創建表格，行數為 JSON 中的項目數量，加上表頭
    table = doc.add_table(rows=len(data) + 1, cols=8)
    table.style = 'Table Grid'

    # 填寫表頭
    headers = [
        "Column 1", "Column 2", "Column 3", "Column 4",
        "Column 5", "Column 6", "Column 7", "Column 8"
    ]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # 填寫每行數據
    for row_idx, row in enumerate(data, start=1):
        for col_idx, (key, value) in enumerate(row.items()):
            table.cell(row_idx, col_idx).text = str(value) if value else ''

# 儲存成 Word 檔案
output_path = 'output0.docx'
doc.save(output_path)

print(f'The table has been created and saved to {output_path}')
