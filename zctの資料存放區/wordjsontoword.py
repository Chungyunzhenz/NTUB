from docx import Document
import json
import math

# 讀取 Word 文件
input_doc_path = "C:/Users/zct/Documents/GitHub/NTUB/zctの資料存放區/output.json"
output_doc_path = 'parsed_output1.docx'
doc = input_doc_path


# 用於存儲處理後的表格數據
output_data = []

# 將段落組合，並嘗試轉換為 JSON
paragraphs = [para.text for para in doc.paragraphs]
full_text = "\n".join(paragraphs)

try:
    # 嘗試解析整個文本為 JSON
    json_data = json.loads(full_text)

    # 提取並匹配文本和座標信息
    pages = json_data.get('pages', [])

    # 定義表格的行數和列數
    num_rows = 20  # 根據需要設置行數
    num_cols = 20  # 根據需要設置列數

    for page in pages:
        page_number = page.get('pageNumber', '')
        dimension = page.get('dimension', {})
        page_width = dimension.get('width', 1)  # 頁面的寬度
        page_height = dimension.get('height', 1)  # 頁面的高度
        blocks = page.get('blocks', [])

        for block in blocks:
            layout = block.get('layout', {})
            text_anchor = layout.get('textAnchor', {})
            bounding_poly = layout.get('boundingPoly', {})
            vertices = bounding_poly.get('vertices', [])

            # 檢查 textSegments 是否存在
            text_segments = text_anchor.get('textSegments', [])
            # 從 JSON 中正確提取字符串
            text_segments_strings = [
                full_text[int(seg.get('startIndex', 0)):int(seg.get('endIndex', 0))]
                for seg in text_segments if 'startIndex' in seg and 'endIndex' in seg
            ]
            block_text = " ".join(text_segments_strings)

            # 根據頂點的坐標計算文字應該放置的表格行和列
            if len(vertices) == 4:
                x_min = min([v.get('x', 0) for v in vertices])
                y_min = min([v.get('y', 0) for v in vertices])
                x_max = max([v.get('x', 0) for v in vertices])
                y_max = max([v.get('y', 0) for v in vertices])

                # 計算相對於頁面的行列位置
                col_min = max(math.floor((x_min / page_width) * num_cols), 0)
                col_max = min(math.ceil((x_max / page_width) * num_cols), num_cols)
                row_min = max(math.floor((y_min / page_height) * num_rows), 0)
                row_max = min(math.ceil((y_max / page_height) * num_rows), num_rows)

                # 將文本與其位置信息儲存
                output_data.append([row_min, row_max, col_min, col_max, block_text])

except json.JSONDecodeError as e:
    print(f"JSON 解析錯誤：{e}")
    print("原始段落內容：")
    for para in paragraphs:
        print(para)

# 創建 Word 文件
output_doc = Document()
table = output_doc.add_table(rows=num_rows, cols=num_cols)

# 初始化每個單元格為空白
for row in table.rows:
    for cell in row.cells:
        cell.text = ''

# 將每個文本片段插入對應的表格單元格
for data in output_data:
    row_min, row_max, col_min, col_max, text = data
    # 檢查矩形區域是否有效
    if row_min < row_max and col_min < col_max:
        try:
            # 在左上角單元格填寫文本
            primary_cell = table.cell(row_min, col_min)
            primary_cell.text = text

            # 合併成矩形的範圍
            for row in range(row_min, row_max):
                for col in range(col_min, col_max):
                    if row == row_min and col == col_min:
                        continue  # 左上角的單元格不需要合併
                    cell_to_merge = table.cell(row, col)
                    primary_cell.merge(cell_to_merge)
        except Exception as e:
            print(f"合併錯誤：{e}")
    else:
        print(f"無效的矩形範圍：row_min={row_min}, row_max={row_max}, col_min={col_min}, col_max={col_max}")

# 保存文件
output_doc.save(output_doc_path)
print(f'已生成新文件：{output_doc_path}')
